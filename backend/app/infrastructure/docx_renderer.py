# Генерация DOCX-резюме

import io
import logging
from typing import Any

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

HSE_BLUE = RGBColor(0x1F, 0x34, 0x69)


def _set_run_font(run, size: int = 11, bold: bool = False, color: RGBColor | None = None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    _set_run_font(run, size=12, bold=True, color=HSE_BLUE)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _add_hyperlink(paragraph, url: str, text: str, color: RGBColor = HSE_BLUE) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Цвет
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "1F3469")
    r_pr.append(c)
    # Подчёркивание
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    new_run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    new_run.append(text_el)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _fetch_image_bytes(url: str) -> bytes | None:
    try:
        with httpx.Client(timeout=5.0) as client:
            r = client.get(url)
            r.raise_for_status()
            return r.content
    except Exception as e:
        logger.warning("Не удалось скачать картинку %s: %s", url, e)
        return None


def render_resume_docx(context: dict[str, Any]) -> bytes:
    owner = context["owner"]
    resume = context["resume"]

    doc = Document()

    photo_url = owner.get("photo_fetch_url") or owner.get("photo_url")
    photo_bytes = _fetch_image_bytes(photo_url) if photo_url else None

    if photo_bytes:
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        left = header_table.cell(0, 0)
        left.width = Inches(2.7)
        try:
            left.paragraphs[0].add_run().add_picture(io.BytesIO(photo_bytes), width=Inches(2.5))
        except Exception as e:
            logger.warning("Не удалось вставить фото: %s", e)
        right = header_table.cell(0, 1)
        name_p = right.paragraphs[0]
        _set_run_font(name_p.add_run(owner["full_name"] or "Без имени"), size=20, bold=True, color=HSE_BLUE)
        if owner.get("email"):
            ep = right.add_paragraph()
            _set_run_font(ep.add_run("Email: "), size=10, bold=True, color=HSE_BLUE)
            _set_run_font(ep.add_run(owner["email"]), size=10, color=RGBColor(0x44, 0x44, 0x44))
        if owner.get("phone"):
            pp = right.add_paragraph()
            _set_run_font(pp.add_run("Телефон: "), size=10, bold=True, color=HSE_BLUE)
            _set_run_font(pp.add_run(owner["phone"]), size=10, color=RGBColor(0x44, 0x44, 0x44))
        date_p = right.add_paragraph()
        _set_run_font(
            date_p.add_run(f"Дата составления: {resume['creation_date']}"),
            size=9, color=RGBColor(0x88, 0x88, 0x88),
        )
    else:
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(name_p.add_run(owner["full_name"] or "Без имени"), size=22, bold=True, color=HSE_BLUE)

        if owner.get("email"):
            ep = doc.add_paragraph()
            _set_run_font(ep.add_run("Email: "), size=10, bold=True, color=HSE_BLUE)
            _set_run_font(ep.add_run(owner["email"]), size=10, color=RGBColor(0x44, 0x44, 0x44))
        if owner.get("phone"):
            pp = doc.add_paragraph()
            _set_run_font(pp.add_run("Телефон: "), size=10, bold=True, color=HSE_BLUE)
            _set_run_font(pp.add_run(owner["phone"]), size=10, color=RGBColor(0x44, 0x44, 0x44))
        date_p = doc.add_paragraph()
        _set_run_font(
            date_p.add_run(f"Дата составления: {resume['creation_date']}"),
            size=9, color=RGBColor(0x88, 0x88, 0x88),
        )

    _add_section_heading(doc, "Целевая должность")
    target_p = doc.add_paragraph()
    _set_run_font(target_p.add_run(resume["position_name"]), size=12, bold=True)
    if resume.get("field_name"):
        sub = f"\nСфера: {resume['field_name']}"
        if resume.get("company"):
            sub += f"  ·  Компания: {resume['company']}"
        _set_run_font(target_p.add_run(sub), size=10, color=RGBColor(0x77, 0x77, 0x77))

    if context["educations"]:
        _add_section_heading(doc, "Образование")
        for edu in context["educations"]:
            p = doc.add_paragraph()
            _set_run_font(
                p.add_run(f"{edu['institution']} — {edu['program_name']}"),
                size=11, bold=True,
            )
            if edu.get("start_year"):
                meta = f"\n{edu['start_year']} — {edu['graduation_year']}"
            else:
                meta = f"\nГод окончания: {edu['graduation_year']}"
            _set_run_font(
                p.add_run(meta),
                size=10, color=RGBColor(0x77, 0x77, 0x77),
            )

    if context["experiences"]:
        _add_section_heading(doc, "Опыт и проекты")
        for exp in context["experiences"]:
            p = doc.add_paragraph()
            _set_run_font(p.add_run(exp["project_name"]), size=11, bold=True)
            meta_parts = [exp["type_name"]]
            if exp.get("discipline_name"):
                meta_parts.append(exp["discipline_name"])
            _set_run_font(
                p.add_run(f"\n{' · '.join(meta_parts)}"),
                size=9, color=RGBColor(0x77, 0x77, 0x77),
            )
            if exp.get("description"):
                desc_p = doc.add_paragraph(exp["description"])
                _set_run_font(desc_p.runs[0], size=10)

            if exp.get("images"):
                attach_p = doc.add_paragraph()
                _set_run_font(
                    attach_p.add_run("Подтверждения: "),
                    size=9, bold=True, color=RGBColor(0x77, 0x77, 0x77),
                )
                for img in exp["images"][:4]:
                    img_bytes = _fetch_image_bytes(img.get("fetch_url") or img["storage_url"])
                    if img_bytes:
                        try:
                            img_p = doc.add_paragraph()
                            img_run = img_p.add_run()
                            img_run.add_picture(io.BytesIO(img_bytes), width=Inches(3.0))
                        except Exception as e:
                            logger.warning("Не удалось вставить картинку: %s", e)
                    link_p = doc.add_paragraph()
                    _set_run_font(
                        link_p.add_run("↗ "),
                        size=10, bold=True, color=HSE_BLUE,
                    )
                    _add_hyperlink(link_p, img["storage_url"], img["file_name"])
                    _set_run_font(
                        link_p.add_run("  — нажмите, чтобы открыть"),
                        size=9, color=RGBColor(0x77, 0x77, 0x77),
                    )

            if exp.get("other_files"):
                link_p = doc.add_paragraph()
                _set_run_font(
                    link_p.add_run("Файлы: "),
                    size=9, bold=True, color=RGBColor(0x77, 0x77, 0x77),
                )
                for i, f in enumerate(exp["other_files"]):
                    if i > 0:
                        _set_run_font(link_p.add_run(" · "), size=9, color=RGBColor(0x99, 0x99, 0x99))
                    _add_hyperlink(link_p, f["storage_url"], f["file_name"])

    if context["skills_by_category"]:
        _add_section_heading(doc, "Навыки")
        for cat, items in context["skills_by_category"].items():
            p = doc.add_paragraph()
            _set_run_font(p.add_run(f"{cat}: "), size=10, bold=True, color=HSE_BLUE)
            _set_run_font(
                p.add_run(", ".join(f"{s['skill_name']} ({s['skill_level']})" for s in items)),
                size=10,
            )

    if context.get("languages"):
        _add_section_heading(doc, "Иностранные языки")
        p = doc.add_paragraph()
        _set_run_font(
            p.add_run(", ".join(f"{l['language_name']} — {l['proficiency']}" for l in context["languages"])),
            size=10,
        )

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        footer_p.add_run("Создано в системе ResumeHelper НИУ ВШЭ"),
        size=8, color=RGBColor(0x99, 0x99, 0x99),
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
